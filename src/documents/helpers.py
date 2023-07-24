from PyPDF2 import PdfReader
import pandas as pd
import docx
import fitz
import re
from io import BytesIO, StringIO
from src.constants import PDF, PLAIN_TEXT, WORD, EXCELL


def handle_pdf(file):
    output = ''
    pdf_reader = PdfReader(file)
    current_page = 0
    while current_page < len(pdf_reader.pages):
        page = pdf_reader.pages[current_page]
        output += page.extract_text()
        current_page += 1
    return output


def handle_docx(file):
    output = ''
    doc = docx.Document(file)
    for para in doc.paragraphs:
        output += para.text
    return output


def handle_excel(file):
    data = pd.read_excel(file)
    return data.to_string(index=False)


def handle_file(file, file_type):
    if file_type == PDF:
        return handle_pdf(file)
    elif file_type == EXCELL:
        return handle_excel(file)
    elif file_type == WORD:
        return handle_docx(file)
    else:
        return ''


def handle_write_pdf(pdf_file, payload):
    pdf_document = fitz.open(pdf_file)
    pages = pdf_document.page_count
    pdf = fitz.open()
    for current_page in range(pages):
        page = pdf_document.load_page(current_page)
        text = page.get_text()
        new_text = text
        expression_to_find = re.findall(r'{{(.*?)}}', new_text)
        for expression in expression_to_find:
            try:
                new_text = re.sub('{{({})}}'.format(
                    expression), payload[expression], new_text)
            except KeyError:
                continue
        pdf.insert_page(-1,
                        text='\n'.join(new_text.splitlines()),
                        fontsize=11,
                        width=595,
                        height=842,
                        fontname='Helvetica',
                        fontfile=None,
                        color=(0, 0, 0))
    output = BytesIO(pdf.tobytes())
    pdf.close()
    pdf_document.close()
    return output


def handle_write_excel(file, payload):
    data = pd.read_excel(file)
    data_as_string = data.to_string(index=False)
    for key in payload.keys():
        data_as_string = re.sub('{{({})}}'.format(
            key), payload[key], data_as_string)
    output_data = StringIO(data_as_string)
    output_data = pd.read_csv(output_data)
    output_data.to_excel('tmp.xlsx')
    output = open('tmp.xlsx', 'rb')
    return output


def handle_write_docx(file, payload):
    output = docx.Document()
    doc = docx.Document(file)
    for para in doc.paragraphs:
        new_text = para.text
        expression_to_find = re.findall(r'{{(.*?)}}', new_text)
        for expression in expression_to_find:
            try:
                new_text = re.sub('{{({})}}'.format(
                    expression), payload[expression], new_text)
            except KeyError:
                continue
        output.add_paragraph(new_text)
    output.save('tmp.docx')
    output_file = open('tmp.docx', 'rb')
    return output_file


def write_file(file, file_type, payload):
    if file_type == PDF:
        return handle_write_pdf(file, payload)
    elif file_type == EXCELL:
        return handle_write_excel(file, payload)
    elif file_type == WORD:
        return handle_write_docx(file, payload)
    else:
        output = ''
        new_text = file
        expression_to_find = re.findall(r'{{(.*?)}}', new_text)
        for expression in expression_to_find:
            try:
                new_text = re.sub('{{({})}}'.format(
                    expression), payload[expression], new_text)
            except KeyError:
                continue
            output += new_text
        return output
